import os
import uuid
import re
from typing import List, Optional, Dict
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from openai import AsyncOpenAI
import docx
from docxtpl import DocxTemplate, RichText
import jinja2
import markdown
from bs4 import BeautifulSoup

app = FastAPI(title="AI Document Constructor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# DeepSeek Configuration
client = AsyncOpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY", "your-deepseek-api-key-here"),
    base_url="https://api.deepseek.com"
)

class EditRequest(BaseModel):
    placeholder_id: str
    user_message: str
    chat_history: Optional[List[Dict[str, str]]] = []
    document_context: Optional[Dict[str, str]] = {}

class ExportRequest(BaseModel):
    placeholder_values: Dict[str, str]

def _process_inline(tag, paragraph):
    """Helper to handle basic formatting inside a paragraph."""
    for child in tag.children:
        if hasattr(child, 'name') and child.name:
            if child.name in ['strong', 'b']:
                paragraph.add_run(child.get_text()).bold = True
            elif child.name in ['em', 'i']:
                paragraph.add_run(child.get_text()).italic = True
            else:
                paragraph.add_run(child.get_text())
        else:
            paragraph.add_run(str(child))

def render_markdown_to_subdoc(md_text, doc):
    """Parses markdown and converts it to Word elements in a SubDoc."""
    sd = doc.new_subdoc()
    # Enable tables extension
    html = markdown.markdown(md_text, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')
    
    for tag in soup.find_all(recursive=False):
        if tag.name == 'p':
            _process_inline(tag, sd.add_paragraph())
        elif tag.name == 'table':
            rows = tag.find_all('tr')
            if not rows: continue
            
            cols_count = len(rows[0].find_all(['td', 'th']))
            table = sd.add_table(rows=0, cols=cols_count)
            table.style = 'Table Grid'
            
            for row in rows:
                cells = row.find_all(['td', 'th'])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    _process_inline(cell, row_cells[i].paragraphs[0])
        elif tag.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            sd.add_heading(tag.get_text(), level=int(tag.name[1]))
        elif tag.name in ['ul', 'ol']:
            for li in tag.find_all('li'):
                p = sd.add_paragraph(style='List Bullet' if tag.name == 'ul' else 'List Number')
                _process_inline(li, p)
                
    return sd

def extract_placeholders_list(docx_path: str):
    doc = docx.Document(docx_path)
    pattern = re.compile(r"<<([^>]+)>>")
    placeholders = []
    seen = set()

    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        for match in matches:
            name = match.split('|')[0].strip()
            if name not in seen:
                placeholders.append({"id": name, "raw": f"<<{match}>>", "constraints": _parse_constraints(match)})
                seen.add(name)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = pattern.findall(para.text)
                    for match in matches:
                        name = match.split('|')[0].strip()
                        if name not in seen:
                            placeholders.append({"id": name, "raw": f"<<{match}>>", "constraints": _parse_constraints(match)})
                            seen.add(name)
    return placeholders

def _parse_constraints(match_str: str) -> Dict[str, str]:
    parts = match_str.split('|')
    constraints = {}
    for part in parts[1:]:
        if ':' in part:
            k, v = part.split(':', 1)
            constraints[k.strip()] = v.strip()
    return constraints

@app.post("/api/v1/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos .docx.")
    
    file_id = f"doc_{uuid.uuid4().hex[:8]}"
    os.makedirs("uploads", exist_ok=True)
    file_path = f"uploads/{file_id}_{file.filename}"
    
    content = await file.read()
    with open(file_path, "wb") as f: f.write(content)
        
    try:
        placeholders = extract_placeholders_list(file_path)
    except Exception as e:
        print(f"Error parsing: {e}")
        placeholders = []

    return {"document_id": file_id, "filename": file.filename, "placeholders": placeholders}

@app.post("/api/v1/documents/{document_id}/edit")
async def edit_placeholder(document_id: str, request: EditRequest):
    try:
        system_content = (
            f"Eres un asistente redactor profesional. Genera el contenido para la sección: {request.placeholder_id}. "
            "IMPORTANTE: Puedes usar Markdown (negritas, cursivas, listas) y si se requiere una tabla, "
            "genérala en formato Markdown (| col | col |). El sistema la convertirá a una tabla de Word real."
        )
        
        if request.document_context:
            context_str = "\n".join([f"- {k}: {v}" for k, v in request.document_context.items() if v])
            if context_str:
                system_content += f"\n\nContexto global:\n{context_str}"

        messages = [{"role": "system", "content": system_content}]
        for msg in request.chat_history:
            messages.append({"role": msg.get("role", "user"), "content": msg.get("content", "")})
        messages.append({"role": "user", "content": request.user_message})

        response = await client.chat.completions.create(
            model="deepseek-chat",
            messages=messages,
            max_tokens=1500,
            temperature=0.7
        )
        
        return {
            "placeholder_id": request.placeholder_id,
            "generated_content": response.choices[0].message.content
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/documents/{document_id}/export")
async def export_document(document_id: str, request: ExportRequest):
    files = [f for f in os.listdir("uploads") if f.startswith(f"{document_id}_")]
    if not files: raise HTTPException(status_code=404, detail="Doc not found")
    
    tpl_path = os.path.join("uploads", files[0])
    doc = DocxTemplate(tpl_path)
    
    # Pre-process tags 
    doc_for_tags = docx.Document(tpl_path)
    pattern = re.compile(r"<<([^>]+)>>")
    full_context = {}
    
    def _collect_tags(paragraphs):
        for para in paragraphs:
            matches = pattern.findall(para.text)
            for m in matches:
                tag_id = m.split('|')[0].strip()
                if tag_id in request.placeholder_values:
                    content = request.placeholder_values[tag_id]
                    # Convert content to SubDoc to handle Rich Text and Tables
                    full_context[m] = render_markdown_to_subdoc(content, doc)

    _collect_tags(doc_for_tags.paragraphs)
    for table in doc_for_tags.tables:
        for row in table.rows:
            for cell in row.cells:
                _collect_tags(cell.paragraphs)

    jinja_env = jinja2.Environment(variable_start_string='<<', variable_end_string='>>')
    doc.render(full_context, jinja_env)
    
    out_path = f"uploads/{document_id}_final.docx"
    doc.save(out_path)
    
    return FileResponse(out_path, filename=f"Propuesta_AI_{files[0].split('_', 1)[1]}", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
